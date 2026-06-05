import json
from fastapi import APIRouter, HTTPException, UploadFile, File
from starlette.responses import StreamingResponse
from ..chat_db import (
    create_conversation, list_conversations,
    get_conversation, delete_conversation, add_message,
    update_conversation_title, undo_messages,
)
from ..memory_service import build_memory_context, extract_memories
from ..agent import agent_chat_stream

router = APIRouter()

# === 对话 API ===
@router.get("/api/conversations")
async def get_conversations_api():
    """获取历史对话列表"""
    conversations = list_conversations()
    return {"conversations": conversations}


@router.get("/api/conversations/{conv_id}")
async def get_conversation_api(conv_id: str):
    """获取单个会话详情"""
    conv = get_conversation(conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="会话不存在")
    return conv


@router.delete("/api/conversations/{conv_id}")
async def delete_conversation_api(conv_id: str):
    """删除会话"""
    if not delete_conversation(conv_id):
        raise HTTPException(status_code=404, detail="会话不存在")
    return {"success": True}


@router.post("/api/conversations/{conv_id}/undo")
async def undo_conversation_api(conv_id: str, body: dict = None):
    """撤回最近 N 轮对话"""
    n = 1
    if body and "n" in body:
        n = max(1, min(int(body["n"]), 10))  # 限制 1-10 轮
    result = undo_messages(conv_id, n)
    if result["undone"] == 0:
        raise HTTPException(status_code=400, detail="没有可撤回的消息")
    return {"success": True, **result}


@router.post("/api/chat")
async def chat_api(body: dict):
    """发送消息并获取 AI 流式响应（SSE）"""
    message = body.get("message", "").strip()
    conv_id = body.get("conversation_id")

    if not message:
        raise HTTPException(status_code=400, detail="消息不能为空")

    # 记录交互（#073 AC5）
    from daozhu.idle_worker import record_interaction
    record_interaction("chat", message[:30])

    # 创建或获取会话
    if not conv_id:
        conv = create_conversation(title=message[:20])
        conv_id = conv["id"]

    # 保存用户消息
    add_message(conv_id, "user", message)

    # 获取会话历史作为上下文
    conv_data = get_conversation(conv_id)
    history = [
        {"role": m["role"], "content": m["content"]}
        for m in conv_data["messages"]
        if m["role"] in ("user", "assistant")  # 过滤掉 tool_call 等非标准 role
    ]

    # 流式响应
    async def generate():
        full_response = ""

        # 构建记忆上下文
        memory_context = build_memory_context(message)

        tool_calls_log = []  # 收集工具调用记录

        async for chunk in agent_chat_stream(history, memory_context=memory_context, conversation_id=conv_id):
            # 工具调用通知（特殊标记）
            if chunk.startswith("[TOOL:"):
                tool_name = chunk[6:-1]
                tool_calls_log.append({"tool": tool_name, "status": "running"})
                yield f"data: {json.dumps({'tool': tool_name, 'conversation_id': conv_id})}\n\n"
                continue
            if chunk.startswith("[TOOL_OK:"):
                tool_name = chunk[9:-1]
                if tool_calls_log and tool_calls_log[-1]["tool"] == tool_name:
                    tool_calls_log[-1]["status"] = "ok"
                yield f"data: {json.dumps({'tool_done': tool_name, 'status': 'ok', 'conversation_id': conv_id})}\n\n"
                continue
            if chunk.startswith("[TOOL_ERR:"):
                parts = chunk[10:-1].split(":", 1)
                tool_name = parts[0]
                err_msg = parts[1] if len(parts) > 1 else ""
                if tool_calls_log and tool_calls_log[-1]["tool"] == tool_name:
                    tool_calls_log[-1]["status"] = "error"
                    tool_calls_log[-1]["error"] = err_msg
                yield f"data: {json.dumps({'tool_done': parts[0], 'status': 'error', 'error': err_msg, 'conversation_id': conv_id})}\n\n"
                continue
            if chunk.startswith("[USAGE:"):
                # Token 使用量统计（#061）
                try:
                    usage_data = json.loads(chunk[7:-1])
                    yield f"data: {json.dumps({'usage': usage_data, 'conversation_id': conv_id})}\n\n"
                except (json.JSONDecodeError, IndexError):
                    pass
                continue
            if chunk == "[COMPACT]":
                yield f"data: {json.dumps({'compact': True, 'conversation_id': conv_id})}\n\n"
                continue

            full_response += chunk
            yield f"data: {json.dumps({'chunk': chunk, 'conversation_id': conv_id})}\n\n"

        # 保存工具调用记录到 conversation（作为 tool_call 类型消息）
        if tool_calls_log:
            add_message(conv_id, "tool_call", json.dumps(tool_calls_log, ensure_ascii=False))

        # 清理 DeepSeek DSML 标记泄露
        import re
        full_response = re.sub(r'</?[|｜]\s*[|｜]?\s*DSML\s*[|｜]\s*[|｜]?[^>]*>', '', full_response)
        full_response = re.sub(r'[|｜]\s*[|｜]?\s*tool_calls\s*>', '', full_response)
        full_response = re.sub(r'[|｜]\s*[|｜]?\s*invoke[^>]*>', '', full_response)
        full_response = re.sub(r'[|｜]\s*[|｜]?\s*parameter[^>]*>', '', full_response)
        full_response = full_response.strip()

        # 保存完整的 AI 回复
        add_message(conv_id, "assistant", full_response)

        # 如果是第一条消息，用内容更新标题
        if len(history) <= 1:
            title = message[:30] + ("..." if len(message) > 30 else "")
            update_conversation_title(conv_id, title)

        # 异步提取记忆（不阻塞响应）
        import asyncio
        all_messages = history + [{"role": "assistant", "content": full_response}]
        asyncio.create_task(extract_memories(all_messages, conv_id))

        yield f"data: {json.dumps({'done': True, 'conversation_id': conv_id})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )



# === 文件上传解析 API ===

ALLOWED_EXTENSIONS = {".docx", ".xlsx", ".xls", ".pdf", ".txt", ".csv"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def _parse_file(filename: str, content: bytes) -> str:
    """解析文件内容为文本"""
    ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式: {ext}。支持: {', '.join(ALLOWED_EXTENSIONS)}")

    if ext == ".txt":
        return content.decode("utf-8", errors="replace")

    if ext == ".csv":
        return content.decode("utf-8", errors="replace")

    if ext == ".docx":
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也读取表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs)

    if ext in (".xlsx", ".xls"):
        import io
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), read_only=False, data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            if len(wb.sheetnames) > 1:
                lines.append(f"[工作表: {sheet_name}]")
            for row in ws.iter_rows(min_row=1, values_only=True):
                cells = [str(c).strip() if c is not None else "" for c in row]
                if any(c for c in cells):  # 跳过全空行
                    lines.append(" | ".join(cells))
        wb.close()
        return "\n".join(lines)

    if ext == ".pdf":
        import io
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)

    raise ValueError(f"解析失败: {ext}")


@router.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """
    上传文件并解析为文本。
    返回解析后的内容，前端可将内容作为消息发给 AI。
    """
    if not file.filename:
        raise HTTPException(400, "缺少文件名")

    # 大小检查
    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(413, f"文件过大（最大 {MAX_FILE_SIZE // 1024 // 1024}MB）")

    try:
        text = _parse_file(file.filename, content)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"文件解析失败: {str(e)[:100]}")

    # 截断过长内容（避免 token 爆炸）
    max_chars = 8000
    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars] + f"\n\n... (内容已截断，共 {len(text)} 字符)"

    return {
        "success": True,
        "filename": file.filename,
        "content": text,
        "chars": len(text),
        "truncated": truncated,
    }
